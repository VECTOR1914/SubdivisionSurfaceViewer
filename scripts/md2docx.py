#!/usr/bin/env python3
"""
md2docx.py — 将 Markdown 实验报告转换为 Word (.docx)
用法: python md2docx.py 实验报告.md 实验报告.docx
"""
import sys, re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_paragraph(doc, text, style='Normal', bold=False, font_size=None, color=None, alignment=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold: run.bold = True
    if font_size: run.font.size = Pt(font_size)
    if color: run.font.color.rgb = RGBColor(*color)
    if alignment is not None: p.alignment = alignment
    return p

def convert(md_path, docx_path):
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    in_table = False
    table_rows = []
    in_bullet = False

    while i < len(lines):
        line = lines[i].rstrip()

        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue
        if in_code_block:
            p = doc.add_paragraph(style='No Spacing')
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(8)
            i += 1
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            # Skip separator lines like |---|---|
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                i += 1
                continue
            cells = [c.strip() for c in line.split('|')[1:-1]]
            table_rows.append(cells)
            # Check if next line is still a table
            if i+1 < len(lines) and '|' in lines[i+1] and lines[i+1].strip().startswith('|'):
                i += 1
                continue
            else:
                # End of table: render it
                if table_rows:
                    table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                    table.style = 'Light Grid Accent 1'
                    for ri, row in enumerate(table_rows):
                        for ci, cell_text in enumerate(row):
                            cell = table.cell(ri, ci)
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(9)
                in_table = False
                table_rows = []
                i += 1
                continue

        # Horizontal rule
        if line.strip() == '---':
            doc.add_paragraph('─' * 60)
            i += 1
            continue

        # Headers
        if line.startswith('# '):
            text = line[2:].strip()
            # Remove LaTeX markup for Word
            add_paragraph(doc, text, bold=True, font_size=18, color=(0,51,102))
            i += 1
            continue
        if line.startswith('## '):
            text = line[3:].strip()
            add_paragraph(doc, text, bold=True, font_size=15, color=(0,70,130))
            i += 1
            continue
        if line.startswith('### '):
            text = line[4:].strip()
            add_paragraph(doc, text, bold=True, font_size=13, color=(0,90,160))
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Image placeholders
        if line.strip().startswith('> ['):
            p = doc.add_paragraph()
            run = p.add_run(line.strip()[2:])
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(180, 100, 0)
            run.italic = True
            # Add a bordered placeholder
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run('[ 图片占位区域 — 请手动插入截图 ]')
            run2.font.size = Pt(11)
            run2.font.color.rgb = RGBColor(150, 150, 150)
            i += 1
            continue

        # Bullet list
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = re.sub(r'^[\s]*[\-\*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            # Clear default text and add formatted
            p.clear()
            # Handle inline code
            parts = re.split(r'(`[^`]+`)', text)
            for part in parts:
                if part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                else:
                    # Handle bold
                    bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
                    for bp in bold_parts:
                        if bp.startswith('**') and bp.endswith('**'):
                            run = p.add_run(bp[2:-2])
                            run.bold = True
                        else:
                            p.add_run(bp)
            i += 1
            continue

        # Regular paragraph
        # Clean up LaTeX math for Word
        text = line.strip()
        text = text.replace('$P_{new}$', 'P_new')
        text = re.sub(r'\$([^$]+)\$', r'\1', text)

        # Handle bold
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        p = doc.add_paragraph()
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                # Handle inline code
                code_parts = re.split(r'(`[^`]+`)', part)
                for cp in code_parts:
                    if cp.startswith('`') and cp.endswith('`'):
                        run = p.add_run(cp[1:-1])
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                    else:
                        p.add_run(cp)
        i += 1

    doc.save(docx_path)
    print(f"已生成: {docx_path}")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print(f"用法: {sys.argv[0]} <输入.md> <输出.docx>")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
